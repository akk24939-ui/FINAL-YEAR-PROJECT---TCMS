"""
generate_doc.py — Smart TASMAC Consumer Regulation System
Crystal-clear diagrams: very large canvas + 200 DPI + big fonts.
"""

import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
from matplotlib import rcParams

# Global font settings — ensures readable text in Word at any zoom
rcParams["font.family"]   = "DejaVu Sans"
rcParams["font.size"]     = 11
rcParams["axes.linewidth"]= 0.6

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_DARK  = "#1F4E79"
C_MID   = "#2E75B6"
C_LITE  = "#D6E4F0"
C_BG    = "#EBF2FA"
C_WHITE = "#FFFFFF"
C_GREY  = "#888888"
C_TEXT  = "#222222"
SAVE_DPI = 200          # 200 DPI on a 24" figure = 4800px wide — crystal clear

# ═════════════════════════════════════════════════════════════════════════════
#  HELPERS
# ═════════════════════════════════════════════════════════════════════════════
def fig_to_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=SAVE_DPI,
                bbox_inches="tight", facecolor=fig.get_facecolor())
    buf.seek(0)
    plt.close(fig)
    return buf

def add_fig(doc, fig, width_cm=16.5, caption=None):
    buf = fig_to_buf(fig)
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r  = cp.runs[0]
        r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(89, 89, 89)

def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#")); tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)

def table_with_header(doc, cols, rows_data, col_widths=None, hdr_color="1F4E79"):
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(cols))
    tbl.style = "Table Grid"
    for i, txt in enumerate(cols):
        c = tbl.rows[0].cells[i]; c.text = ""
        shade_cell(c, hdr_color)
        r = c.paragraphs[0].add_run(txt)
        r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            c = tbl.rows[ri + 1].cells[ci]; c.text = val
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        set_col_widths(tbl, col_widths)
    return tbl

# ── TOC ───────────────────────────────────────────────────────────────────────
TOC_ENTRIES = [
    (1, "1. Project Overview",                   3),
    (1, "2. Database Design — Entity Tables",    4),
    (2, "    2.1  Users",                        4),
    (2, "    2.2  Roles",                        4),
    (2, "    2.3  Consumers",                    4),
    (2, "    2.4  Shops",                        5),
    (2, "    2.5  Products",                     5),
    (2, "    2.6  Purchases",                    5),
    (2, "    2.7  Restrictions",                 6),
    (2, "    2.8  QR Codes",                     6),
    (2, "    2.9  Notifications",                6),
    (2, "    2.10 Summary of Relationships",     7),
    (1, "3. Entity-Relationship (ER) Diagram",   8),
    (1, "4. Data Flow Diagrams (DFD)",           9),
    (2, "    4.1  Context-Level DFD (Level 0)",  9),
    (2, "    4.2  Level-1 DFD",                 10),
    (2, "    4.3  Processes",                   11),
    (2, "    4.4  Data Stores",                 11),
    (1, "5. Use Case Design",                   12),
    (2, "    5.1  Actors",                      12),
    (2, "    5.2  Use Case Descriptions",       13),
]

def add_toc(doc):
    for lvl, title, pg in TOC_ENTRIES:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        pPr  = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab  = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right"); tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab); pPr.append(tabs)
        r1 = p.add_run(title); r1.font.size = Pt(10)
        if lvl == 1:
            r1.bold = True; r1.font.color.rgb = RGBColor(31, 78, 121)
        else:
            r1.font.color.rgb = RGBColor(46, 117, 182)
        r2 = p.add_run(f"\t{pg}")
        r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(89, 89, 89)

# ═════════════════════════════════════════════════════════════════════════════
#  DIAGRAM 1 — ER DIAGRAM
#  Canvas: 32" × 44"  →  6400 × 8800 px at 200 DPI — very crisp in Word
# ═════════════════════════════════════════════════════════════════════════════
def draw_er_diagram():
    FW, FH = 32, 44               # inches — huge canvas for clarity
    fig, ax = plt.subplots(figsize=(FW, FH))
    ax.set_xlim(0, FW); ax.set_ylim(0, FH)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FBFF")

    BOX_W = 8.5
    HDR_H = 1.1
    ROW_H = 0.78

    # ── Entity definitions ──────────────────────────────────────────────────
    # Each: (left_x, bottom_y, "TITLE", [(lbl, field), ...])
    entities = [
        # COLUMN A  x=0.5
        (0.5, 35.5, "USERS", [
            ("PK", "id (UUID)"),
            ("",   "email"),
            ("",   "password_hash"),
            ("FK", "role_id"),
            ("",   "mfa_enabled"),
            ("",   "is_active"),
            ("",   "created_at"),
        ]),
        (0.5, 22.0, "AUDIT_LOGS", [
            ("PK", "id (UUID)"),
            ("FK", "actor_user_id"),
            ("",   "action"),
            ("",   "target_table"),
            ("",   "target_id"),
            ("",   "ip_address"),
            ("",   "timestamp"),
        ]),
        (0.5, 9.0, "NOTIFICATIONS", [
            ("PK", "id (UUID)"),
            ("FK", "user_id"),
            ("",   "type"),
            ("",   "message"),
            ("",   "read_flag"),
            ("",   "created_at"),
        ]),

        # COLUMN B  x=11.5
        (11.5, 39.0, "ROLES", [
            ("PK", "id (UUID)"),
            ("",   "name"),
        ]),
        (11.5, 29.0, "CONSENTS", [
            ("PK", "id (UUID)"),
            ("FK", "consumer_id"),
            ("FK", "caretaker_user_id"),
            ("",   "scope"),
            ("",   "granted_at"),
            ("",   "revoked_at"),
        ]),
        (11.5, 18.5, "QR_CODES", [
            ("PK", "id (UUID)"),
            ("FK", "consumer_id"),
            ("",   "signed_token"),
            ("",   "issued_at"),
            ("",   "expires_at"),
            ("",   "used_flag"),
        ]),
        (11.5, 7.5, "RESTRICTIONS", [
            ("PK", "id (UUID)"),
            ("FK", "consumer_id"),
            ("",   "daily_limit"),
            ("",   "weekly_limit"),
            ("",   "monthly_limit"),
            ("",   "self_restricted"),
            ("",   "effective_from"),
        ]),

        # COLUMN C  x=22.5
        (22.5, 39.0, "CONSUMERS", [
            ("PK", "id (UUID)"),
            ("FK", "user_id"),
            ("",   "mock_id_number_enc"),
            ("",   "dob"),
            ("",   "gender"),
            ("",   "district"),
            ("",   "teetotaler_flag"),
        ]),
        (22.5, 28.5, "PURCHASES", [
            ("PK", "id (UUID)"),
            ("FK", "consumer_id"),
            ("FK", "shop_id"),
            ("FK", "product_id"),
            ("",   "quantity"),
            ("",   "timestamp"),
            ("",   "idempotency_key"),
        ]),
        (22.5, 18.5, "PRODUCTS", [
            ("PK", "id (UUID)"),
            ("",   "name"),
            ("",   "category"),
            ("",   "volume_ml"),
            ("",   "standard_drink_equiv"),
            ("",   "price"),
        ]),
        (22.5, 9.5, "SHOPS", [
            ("PK", "id (UUID)"),
            ("",   "name"),
            ("",   "district"),
            ("",   "license_no"),
            ("FK", "operator_user_id"),
        ]),
        (22.5, 0.8, "REPORTS", [
            ("PK", "id (UUID)"),
            ("FK", "generated_by"),
            ("",   "type"),
            ("",   "scope"),
            ("",   "file_path"),
            ("",   "created_at"),
        ]),
    ]

    centers = {}

    for bx, by, title, fields in entities:
        nf     = len(fields)
        body_h = nf * ROW_H

        # Header rectangle
        ax.add_patch(FancyBboxPatch(
            (bx, by + body_h), BOX_W, HDR_H,
            boxstyle="square,pad=0", lw=2,
            edgecolor=C_DARK, facecolor=C_DARK, zorder=4))
        ax.text(bx + BOX_W / 2, by + body_h + HDR_H / 2,
                title, ha="center", va="center",
                fontsize=15, fontweight="bold", color="white", zorder=6)

        # Body rectangle
        ax.add_patch(FancyBboxPatch(
            (bx, by), BOX_W, body_h,
            boxstyle="square,pad=0", lw=2,
            edgecolor=C_MID, facecolor=C_WHITE, zorder=4))

        # Field rows
        for fi, (lbl, fname) in enumerate(fields):
            ry  = by + (nf - fi - 1) * ROW_H
            # alternating background
            if fi % 2 == 1:
                ax.add_patch(FancyBboxPatch(
                    (bx + 0.05, ry + 0.04), BOX_W - 0.10, ROW_H - 0.08,
                    boxstyle="square,pad=0", lw=0,
                    facecolor="#EBF2FA", zorder=4))
            color  = C_MID  if lbl else C_TEXT
            weight = "bold" if lbl == "PK" else "normal"
            label  = f"  {lbl}   {fname}" if lbl else f"      {fname}"
            ax.text(bx + 0.25, ry + ROW_H / 2,
                    label, ha="left", va="center",
                    fontsize=12, fontweight=weight, color=color, zorder=6)
            # row divider
            if fi < nf - 1:
                ax.plot([bx, bx + BOX_W], [ry, ry],
                        color="#C8DCF0", lw=0.8, zorder=5)

        # Center point for relationship lines
        cx = bx + BOX_W / 2
        cy = by + (body_h + HDR_H) / 2
        centers[title] = (cx, cy)

    # ── Relationships ───────────────────────────────────────────────────────
    rels = [
        ("ROLES",     "USERS",        "1 : N"),
        ("USERS",     "CONSUMERS",    "1 : 1"),
        ("USERS",     "SHOPS",        "1 : N"),
        ("USERS",     "AUDIT_LOGS",   "1 : N"),
        ("USERS",     "NOTIFICATIONS","1 : N"),
        ("USERS",     "REPORTS",      "1 : N"),
        ("CONSUMERS", "PURCHASES",    "1 : N"),
        ("CONSUMERS", "RESTRICTIONS", "1 : N"),
        ("CONSUMERS", "QR_CODES",     "1 : N"),
        ("CONSUMERS", "CONSENTS",     "1 : N"),
        ("SHOPS",     "PURCHASES",    "1 : N"),
        ("PRODUCTS",  "PURCHASES",    "1 : N"),
    ]
    for fr, to, lbl in rels:
        if fr not in centers or to not in centers:
            continue
        x1, y1 = centers[fr]
        x2, y2 = centers[to]
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-",
                                    color="#AABBCC", lw=1.8,
                                    connectionstyle="arc3,rad=0.08"),
                    zorder=2)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my, f" {lbl} ", ha="center", va="center",
                fontsize=11, fontweight="bold", color=C_DARK, zorder=7,
                bbox=dict(boxstyle="round,pad=0.22",
                          fc="white", ec=C_MID, lw=1.2, alpha=0.95))

    ax.set_title("ER Diagram — Smart TASMAC Consumer Regulation System",
                 fontsize=22, fontweight="bold", color=C_DARK, pad=20)
    fig.tight_layout(pad=1.5)
    return fig


# ═════════════════════════════════════════════════════════════════════════════
#  DIAGRAM 2 — DFD LEVEL 0  (Context Diagram)
#  Canvas: 22" × 14"  →  4400 × 2800 px at 200 DPI
# ═════════════════════════════════════════════════════════════════════════════
def draw_dfd_level0():
    FW, FH = 22, 14
    fig, ax = plt.subplots(figsize=(FW, FH))
    ax.set_xlim(0, FW); ax.set_ylim(0, FH)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FBFF")

    # Central circle
    cx, cy, R = 11.0, 7.0, 3.2
    circle = plt.Circle((cx, cy), R, color=C_DARK, zorder=4)
    ax.add_patch(circle)
    for txt, dy in [("Smart TASMAC", 0.85),
                    ("Consumer Regulation", 0.0),
                    ("System", -0.85)]:
        ax.text(cx, cy + dy, txt, ha="center", va="center",
                fontsize=17, fontweight="bold", color="white", zorder=6)

    AW, AH = 3.5, 1.3
    actors = [
        # (bx, cy, label, port_side)
        (0.3,  12.0, "CONSUMER",          "right"),
        (0.3,   8.5, "SHOP OPERATOR",     "right"),
        (0.3,   4.5, "GOVERNMENT ADMIN",  "right"),
        (0.3,   1.2, "CARETAKER",         "right"),
        (18.2,  6.5, "DOCTOR",            "left"),
    ]

    ports = {}
    for bx, ay, lbl, side in actors:
        ax.add_patch(FancyBboxPatch(
            (bx, ay - AH / 2), AW, AH,
            boxstyle="square,pad=0.12", lw=2.2,
            edgecolor=C_DARK, facecolor=C_BG, zorder=5))
        ax.text(bx + AW / 2, ay, lbl,
                ha="center", va="center",
                fontsize=14, fontweight="bold", color=C_DARK, zorder=6)
        ports[lbl] = (bx + AW, ay) if side == "right" else (bx, ay)

    # ── Data flows ──
    flows = [
        ("CONSUMER",
         cx - R * 0.82, cy + R * 0.54,
         "Register · Request QR · Set Limits",
         "QR Code · Confirmation · Alerts"),
        ("SHOP OPERATOR",
         cx - R * 0.95, cy + R * 0.08,
         "Scan QR · Record Purchase",
         "Eligibility · Receipt"),
        ("GOVERNMENT ADMIN",
         cx - R * 0.82, cy - R * 0.54,
         "Request Analytics",
         "District / Revenue Reports"),
        ("CARETAKER",
         cx - R * 0.55, cy - R * 0.84,
         "Request Status (consent-gated)",
         "Health Alerts"),
        ("DOCTOR",
         cx + R * 0.95, cy,
         "Request Health Data",
         "Anonymised Health Trends"),
    ]

    for albl, ex, ey, inp, out in flows:
        sx, sy = ports[albl]
        # Input arrow
        ax.annotate("", xy=(ex, ey + 0.25), xytext=(sx, sy),
                    arrowprops=dict(arrowstyle="->",
                                    color=C_MID, lw=2.2), zorder=3)
        # Output arrow
        ax.annotate("", xy=(sx, sy), xytext=(ex, ey - 0.25),
                    arrowprops=dict(arrowstyle="->",
                                    color=C_DARK, lw=1.8,
                                    linestyle="dashed"), zorder=3)
        mx = (sx + ex) / 2; my = (sy + ey) / 2
        ax.text(mx, my + 0.35, inp, ha="center", va="bottom",
                fontsize=11, color=C_DARK,
                bbox=dict(boxstyle="round,pad=0.22",
                          fc="white", ec=C_MID, lw=1, alpha=0.95))
        ax.text(mx, my - 0.35, out, ha="center", va="top",
                fontsize=10.5, color="#555", style="italic",
                bbox=dict(boxstyle="round,pad=0.22",
                          fc="white", ec="#AAAAAA", lw=0.8, alpha=0.95))

    ax.set_title("Figure 4.1 — Context-Level DFD (Level 0)",
                 fontsize=16, style="italic", color="#555", pad=14)
    fig.tight_layout(pad=1.5)
    return fig


# ═════════════════════════════════════════════════════════════════════════════
#  DIAGRAM 3 — DFD LEVEL 1
#  Canvas: 26" × 18"  →  5200 × 3600 px at 200 DPI
# ═════════════════════════════════════════════════════════════════════════════
def draw_dfd_level1():
    FW, FH = 26, 20
    fig, ax = plt.subplots(figsize=(FW, FH))
    ax.set_xlim(0, FW); ax.set_ylim(0, FH)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FBFF")

    BPW, BPH = 6.5, 1.20
    SDW, SDH = 6.5, 0.80
    AW,  AH  = 3.5, 1.00

    # (x, cy, id, label)
    processes = [
        (6.5, 18.2, "1.0", "Auth & RBAC"),
        (6.5, 15.6, "2.0", "Profile & Restriction Mgmt"),
        (6.5, 13.0, "3.0", "QR Issuance & Verification"),
        (6.5, 10.4, "4.0", "Purchase Recording"),
        (6.5,  7.8, "5.0", "Analytics & Reporting"),
        (6.5,  5.2, "6.0", "Notification Engine"),
    ]

    # (sx, cy, label)
    stores = [
        (16.5, 18.2, "D1 — users / roles"),
        (16.5, 15.6, "D2 — consumers / restrictions / consents"),
        (16.5, 13.0, "D3 — qr_codes"),
        (16.5, 10.4, "D4 — purchases / products / shops"),
        (16.5,  7.8, "D5 — reports"),
        (16.5,  5.2, "D6 — notifications / audit_logs"),
    ]

    # actors on left
    actor_lbl = ["Consumer", "Consumer", "Consumer /\nShop Operator",
                  "Shop Operator", "Gov Admin / Doctor", "All Actors"]
    for i, (_, py, _, plbl) in enumerate(processes):
        albl = actor_lbl[i]
        ax.add_patch(FancyBboxPatch(
            (0.3, py - AH / 2), AW, AH,
            boxstyle="square,pad=0.08", lw=1.5,
            edgecolor="#777", facecolor="#F2F2F2", zorder=3))
        ax.text(0.3 + AW / 2, py, albl,
                ha="center", va="center",
                fontsize=11.5, color="#333", zorder=4)
        ax.annotate("", xy=(6.5, py), xytext=(0.3 + AW, py),
                    arrowprops=dict(arrowstyle="->", color=C_MID, lw=1.8), zorder=2)

    # Draw processes
    for px, py, pid, plbl in processes:
        ax.add_patch(FancyBboxPatch(
            (px, py - BPH / 2), BPW, BPH,
            boxstyle="round,pad=0.12", lw=2,
            edgecolor=C_MID, facecolor=C_LITE, zorder=4))
        ax.text(px + 0.55, py + 0.15, pid,
                ha="left", va="center",
                fontsize=11, fontweight="bold", color=C_DARK, zorder=5)
        ax.text(px + BPW / 2, py - 0.13, plbl,
                ha="center", va="center",
                fontsize=13, color=C_DARK, zorder=5)
        # Arrow to data store
        ax.annotate("", xy=(16.5, py), xytext=(px + BPW, py),
                    arrowprops=dict(arrowstyle="<->", color="#777", lw=1.5), zorder=2)

    # Draw data stores (open-ended symbol)
    for sx, sy, slbl in stores:
        ax.plot([sx, sx + SDW], [sy + SDH / 2, sy + SDH / 2],
                color=C_DARK, lw=2.2, zorder=3)
        ax.plot([sx, sx + SDW], [sy - SDH / 2, sy - SDH / 2],
                color=C_DARK, lw=2.2, zorder=3)
        ax.plot([sx, sx], [sy - SDH / 2, sy + SDH / 2],
                color=C_DARK, lw=2.2, zorder=3)
        ax.add_patch(FancyBboxPatch(
            (sx, sy - SDH / 2), SDW, SDH,
            boxstyle="square,pad=0", lw=0,
            facecolor=C_BG, zorder=2))
        ax.text(sx + 0.30, sy, slbl,
                ha="left", va="center",
                fontsize=12, color=C_DARK, zorder=4)

    ax.set_title("Figure 4.2 — Level-1 DFD",
                 fontsize=16, style="italic", color="#555", pad=14)
    fig.tight_layout(pad=1.5)
    return fig


# ═════════════════════════════════════════════════════════════════════════════
#  DIAGRAM 4 — USE CASE DIAGRAM
#  Canvas: 22" × 26"  →  4400 × 5200 px at 200 DPI
# ═════════════════════════════════════════════════════════════════════════════
def draw_use_case():
    FW, FH = 22, 26
    fig, ax = plt.subplots(figsize=(FW, FH))
    ax.set_xlim(0, FW); ax.set_ylim(0, FH)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FBFF")

    # System boundary
    ax.add_patch(FancyBboxPatch(
        (5.5, 0.5), 15.8, 24.8,
        boxstyle="square,pad=0.1", lw=2.5,
        edgecolor=C_MID, facecolor="#F4F8FD", zorder=1))
    ax.text(13.4, 25.1,
            "Smart TASMAC Consumer Regulation System",
            ha="center", va="center",
            fontsize=16, fontweight="bold", color=C_DARK)

    # Use cases (x_centre, y_centre, label)
    uc_yw = 1.8        # spacing between ovals
    use_cases = [
        (13.4, 23.5, "Register / Login"),
        (13.4, 21.7, "Manage Profile"),
        (13.4, 19.9, "Set Consumption Limits"),
        (13.4, 18.1, "Enable Self-Restriction / Teetotaler"),
        (13.4, 16.3, "View Purchase History"),
        (13.4, 14.5, "Download / Request PDF Report"),
        (13.4, 12.7, "Generate / View QR Profile"),
        (13.4, 10.9, "Scan QR & Verify Consumer"),
        (13.4,  9.1, "Record Purchase"),
        (13.4,  7.3, "View Remaining Limit / Receipt"),
        (13.4,  5.5, "View District Analytics / Export"),
        (13.4,  3.7, "View Anonymised Health Trends"),
        (13.4,  1.9, "Grant / Revoke / View Consent"),
    ]

    EW, EH = 9.0, 1.10
    for ux, uy, lbl in use_cases:
        ax.add_patch(mpatches.Ellipse(
            (ux, uy), EW, EH,
            edgecolor=C_MID, facecolor=C_LITE,
            lw=2.0, zorder=3))
        ax.text(ux, uy, lbl, ha="center", va="center",
                fontsize=13, fontweight="bold", color=C_DARK, zorder=4)

    # Actors: (x_centre, y_centre, label, list_of_uc_y_to_connect)
    AW2, AH2 = 3.0, 0.95
    actors = [
        (2.2, 17.2, "Consumer",
         [23.5, 21.7, 19.9, 18.1, 16.3, 14.5, 12.7, 1.9]),
        (2.2, 10.0, "Shop\nOperator",
         [10.9, 9.1, 7.3]),
        (2.2,  5.5, "Government\nAdmin",
         [5.5]),
        (2.2,  3.7, "Doctor",
         [3.7]),
        (2.2,  1.9, "Caretaker",
         [1.9]),
    ]

    for ax_, ay_, lbl, uc_ys in actors:
        ax.add_patch(FancyBboxPatch(
            (ax_ - AW2 / 2, ay_ - AH2 / 2), AW2, AH2,
            boxstyle="square,pad=0.10", lw=2,
            edgecolor=C_DARK, facecolor=C_BG, zorder=4))
        ax.text(ax_, ay_, lbl, ha="center", va="center",
                fontsize=13, fontweight="bold", color=C_DARK, zorder=5)
        for uc_y in uc_ys:
            ax.plot([ax_ + AW2 / 2, 8.95], [ay_, uc_y],
                    color="#AAAACC", lw=1.4, zorder=2)

    ax.set_title("Figure 5.1 — Use Case Diagram",
                 fontsize=16, style="italic", color="#555", pad=12)
    fig.tight_layout(pad=1.5)
    return fig


# ═════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Cover
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\nSMART TASMAC CONSUMER REGULATION SYSTEM")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(31, 78, 121)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("\nSystem Design Document\nTables  •  ER Diagram  •  DFD  •  Use Case Design\n\n2025 – 2026")
r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(46, 117, 182)
doc.add_page_break()

# TOC
h_toc = doc.add_heading("Table of Contents", level=1)
h_toc.runs[0].font.color.rgb = RGBColor(31, 78, 121)
add_toc(doc)
doc.add_page_break()

# Section 1
h1 = doc.add_heading("1. Project Overview", level=1)
h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
doc.add_paragraph(
    "The Smart TASMAC Consumer Regulation System is a security-hardened web application "
    "promoting responsible alcohol consumption through consumer profiles, purchase tracking, "
    "self-imposed limits, and administrative dashboards. Five roles are supported: "
    "Consumer, Shop Operator, Government Administrator, Doctor, and Caretaker. "
    "All identity verification uses simulated (synthetic) data — no real Aadhaar or UIDAI integration."
).runs[0].font.size = Pt(10)
doc.add_page_break()

# Section 2
h2 = doc.add_heading("2. Database Design — Entity Tables", level=1)
h2.runs[0].font.color.rgb = RGBColor(31, 78, 121)

tables_data = [
    ("2.1  users", "Central authentication identity store for every user.", [
        ("id","UUID","PRIMARY KEY"),("email","VARCHAR(255)","UNIQUE, NOT NULL"),
        ("password_hash","TEXT","Argon2id hash — NOT NULL"),("role_id","UUID","FK → roles.id"),
        ("mfa_enabled","BOOLEAN","DEFAULT FALSE"),("is_active","BOOLEAN","DEFAULT TRUE"),
        ("created_at","TIMESTAMPTZ","DEFAULT now()")]),
    ("2.2  roles", "Five RBAC role definitions.", [
        ("id","UUID","PRIMARY KEY"),
        ("name","VARCHAR(50)","consumer | shop_operator | gov_admin | doctor | caretaker")]),
    ("2.3  consumers", "Extended consumer profile. mock_id_number is AES-GCM encrypted.", [
        ("id","UUID","PRIMARY KEY"),("user_id","UUID","FK → users.id, UNIQUE"),
        ("mock_id_number_enc","TEXT","AES-GCM encrypted simulated ID"),("dob","DATE","NOT NULL"),
        ("gender","VARCHAR(10)","CHECK IN ('M','F','Other')"),("district","VARCHAR(100)","NOT NULL"),
        ("teetotaler_flag","BOOLEAN","DEFAULT FALSE")]),
    ("2.4  shops", "TASMAC retail outlet records.", [
        ("id","UUID","PRIMARY KEY"),("name","VARCHAR(200)","NOT NULL"),
        ("district","VARCHAR(100)","NOT NULL"),("license_no","VARCHAR(50)","UNIQUE, NOT NULL"),
        ("operator_user_id","UUID","FK → users.id")]),
    ("2.5  products", "Alcohol product catalogue.", [
        ("id","UUID","PRIMARY KEY"),("name","VARCHAR(200)","NOT NULL"),
        ("category","VARCHAR(50)","Beer | Spirits | Wine"),("volume_ml","INTEGER","NOT NULL"),
        ("standard_drink_equiv","NUMERIC(5,2)","Calculated"),("price","NUMERIC(10,2)","NOT NULL")]),
    ("2.6  purchases", "Immutable purchase ledger. Idempotency key prevents double-recording.", [
        ("id","UUID","PRIMARY KEY"),("consumer_id","UUID","FK → consumers.id"),
        ("shop_id","UUID","FK → shops.id"),("product_id","UUID","FK → products.id"),
        ("quantity","INTEGER","CHECK > 0"),("timestamp","TIMESTAMPTZ","DEFAULT now()"),
        ("idempotency_key","VARCHAR(64)","UNIQUE, NOT NULL")]),
    ("2.7  restrictions", "Consumer-set consumption limits per time window.", [
        ("id","UUID","PRIMARY KEY"),("consumer_id","UUID","FK → consumers.id"),
        ("daily_limit","NUMERIC(5,2)","Standard drinks / day"),
        ("weekly_limit","NUMERIC(5,2)","Standard drinks / week"),
        ("monthly_limit","NUMERIC(5,2)","Standard drinks / month"),
        ("self_restricted","BOOLEAN","DEFAULT FALSE"),("effective_from","DATE","NOT NULL")]),
    ("2.8  qr_codes", "Signed, time-limited, single-use QR tokens.", [
        ("id","UUID","PRIMARY KEY"),("consumer_id","UUID","FK → consumers.id"),
        ("signed_token","TEXT","HMAC-SHA256 opaque token, UNIQUE"),
        ("issued_at","TIMESTAMPTZ","DEFAULT now()"),("expires_at","TIMESTAMPTZ","15-minute TTL"),
        ("used_flag","BOOLEAN","DEFAULT FALSE — single-use enforcement")]),
    ("2.9  notifications", "In-app notification messages.", [
        ("id","UUID","PRIMARY KEY"),("user_id","UUID","FK → users.id"),
        ("type","VARCHAR(50)","LIMIT_ALERT | CONSENT_REQUEST | SYSTEM"),
        ("message","TEXT","NOT NULL"),("read_flag","BOOLEAN","DEFAULT FALSE"),
        ("created_at","TIMESTAMPTZ","DEFAULT now()")]),
    ("2.10  Summary of Relationships", None, [
        ("ROLES","1 : N","→ USERS"),("USERS","1 : 1","→ CONSUMERS"),
        ("USERS","1 : N","→ SHOPS (as operator)"),("CONSUMERS","1 : N","→ PURCHASES"),
        ("CONSUMERS","1 : N","→ RESTRICTIONS"),("CONSUMERS","1 : N","→ QR_CODES"),
        ("CONSUMERS","N : M","→ USERS (caretaker) via CONSENTS"),
        ("SHOPS","1 : N","→ PURCHASES"),("PRODUCTS","1 : N","→ PURCHASES"),
        ("USERS","1 : N","→ AUDIT_LOGS"),("USERS","1 : N","→ REPORTS")]),
]

for title, desc, rows in tables_data:
    h = doc.add_heading(title, level=2)
    h.runs[0].font.color.rgb = RGBColor(46, 117, 182)
    if desc:
        dp = doc.add_paragraph(desc)
        dp.runs[0].italic = True; dp.runs[0].font.size = Pt(9)
        table_with_header(doc, ["Column","Data Type","Constraints / Notes"],
                          rows, col_widths=[4.5, 3.5, 9])
    else:
        table_with_header(doc, ["From Entity","Cardinality","To Entity"],
                          rows, col_widths=[4, 3, 10.5])
    doc.add_paragraph()

doc.add_page_break()

# Section 3 — ER
h3 = doc.add_heading("3. Entity-Relationship (ER) Diagram", level=1)
h3.runs[0].font.color.rgb = RGBColor(31, 78, 121)
doc.add_paragraph(
    "All 12 entities with primary keys (PK), foreign keys (FK), and relationship cardinality."
).runs[0].font.size = Pt(10)
print("Rendering ER diagram...")
add_fig(doc, draw_er_diagram(), width_cm=16.5, caption="Figure 3.1 — ER Diagram")
doc.add_page_break()

# Section 4 — DFD
h4 = doc.add_heading("4. Data Flow Diagrams (DFD)", level=1)
h4.runs[0].font.color.rgb = RGBColor(31, 78, 121)
doc.add_heading("4.1  Context-Level DFD (Level 0)", level=2).runs[0].font.color.rgb = RGBColor(46,117,182)
doc.add_paragraph(
    "The system as a single process interacting with five external entities."
).runs[0].font.size = Pt(10)
print("Rendering DFD Level 0...")
add_fig(doc, draw_dfd_level0(), width_cm=16.5, caption="Figure 4.1 — Context-Level DFD (Level 0)")
doc.add_paragraph()

doc.add_heading("4.2  Level-1 DFD", level=2).runs[0].font.color.rgb = RGBColor(46,117,182)
doc.add_paragraph("Six processes, six data stores.").runs[0].font.size = Pt(10)
print("Rendering DFD Level 1...")
add_fig(doc, draw_dfd_level1(), width_cm=16.5, caption="Figure 4.2 — Level-1 DFD")
doc.add_paragraph()

doc.add_heading("4.3  Processes", level=2).runs[0].font.color.rgb = RGBColor(46,117,182)
table_with_header(doc, ["ID","Process Name","Data Stores"],
    [("1.0","Auth & RBAC","users, roles"),
     ("2.0","Profile & Restriction Mgmt","consumers, restrictions, consents"),
     ("3.0","QR Issuance & Verification","qr_codes, consumers"),
     ("4.0","Purchase Recording","purchases, products, shops"),
     ("5.0","Analytics & Reporting","reports, purchases"),
     ("6.0","Notification Engine","notifications, audit_logs")],
    col_widths=[2, 6, 9.5])
doc.add_paragraph()

doc.add_heading("4.4  Data Stores", level=2).runs[0].font.color.rgb = RGBColor(46,117,182)
table_with_header(doc, ["Store","Table(s)","Description"],
    [("D1","users, roles","Auth & role data"),
     ("D2","consumers, restrictions","Consumer profiles & limits"),
     ("D3","qr_codes","Signed, time-limited QR tokens"),
     ("D4","purchases, products, shops","Transaction ledger & catalogue"),
     ("D5","reports","Generated report metadata"),
     ("D6","notifications, audit_logs","Alerts & security event trail")],
    col_widths=[2, 5, 10.5])
doc.add_page_break()

# Section 5 — Use Case
h5 = doc.add_heading("5. Use Case Design", level=1)
h5.runs[0].font.color.rgb = RGBColor(31, 78, 121)
doc.add_heading("5.1  Actors", level=2).runs[0].font.color.rgb = RGBColor(46,117,182)
table_with_header(doc, ["Actor","Description"],
    [("Consumer","Manages profile, limits, QR, purchase history"),
     ("Shop Operator","Scans QR, records purchases, prints receipts"),
     ("Government Admin","Views aggregated analytics, exports reports"),
     ("Doctor","Views anonymised health trends & risk data"),
     ("Caretaker","Views consent-gated status, receives alerts")],
    col_widths=[4, 13.5])
doc.add_paragraph()
print("Rendering Use Case diagram...")
add_fig(doc, draw_use_case(), width_cm=15, caption="Figure 5.1 — Use Case Diagram")
doc.add_paragraph()

doc.add_heading("5.2  Use Case Descriptions", level=2).runs[0].font.color.rgb = RGBColor(46,117,182)
table_with_header(doc, ["ID","Use Case","Actor","Description"],
    [("UC-01","Register Account","Consumer","Email, password, DOB, district, mock-ID"),
     ("UC-02","Login","Consumer","JWT (15 min) + httpOnly refresh cookie"),
     ("UC-03","Set Consumption Limits","Consumer","Daily/weekly/monthly standard-drink limits"),
     ("UC-04","Enable Self-Restriction","Consumer","Block all purchases; teetotaler toggle"),
     ("UC-05","View Purchase History","Consumer","Paginated, filterable purchase list"),
     ("UC-06","Download PDF Report","Consumer","Personal consumption report"),
     ("UC-07","Generate / View QR","Consumer","HMAC-signed 15-min single-use QR code"),
     ("UC-08","Grant / Revoke Consent","Consumer","Create or revoke caretaker access consent"),
     ("UC-09","Scan QR & Verify","Shop Operator","Validate QR signature, TTL, used_flag"),
     ("UC-10","Record Purchase","Shop Operator","Idempotency check + limit enforcement"),
     ("UC-11","View Remaining Limit","Shop Operator","Live limit status after purchase"),
     ("UC-12","View Analytics","Gov Admin","Aggregated stats — k-anonymity >= 5"),
     ("UC-13","Export Report","Gov Admin","PDF / Excel / JSON — rate-limited"),
     ("UC-14","View Health Trends","Doctor","Anonymous age-group consumption charts"),
     ("UC-15","View Risk Data","Doctor","Risk-tier aggregates, no PII"),
     ("UC-16","View Consumer Status","Caretaker","Consent-gated — 403 if no consent"),
     ("UC-17","Receive Alerts","Caretaker","Notifications on limit breach")],
    col_widths=[1.7, 4.5, 3.5, 7.8])

# Save
out = r"A:\FINAL YEAR PROJECT TASMAC\TASMAC_Design_NEW.docx"
doc.save(out)
print("Saved -> " + out)

